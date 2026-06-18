"""
生成 A9 智能家居设备控制系统 — 项目框架文档 (Word)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_code_block(doc, code_text):
    """添加代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(line, style='No Spacing')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    doc.add_paragraph()

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 1.0)
    return p

# ══════════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第十五届"中国软件杯"大学生软件设计大赛')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('A9 — 基于OpenHarmony的智能家居设备控制系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('项 目 框 架 文 档')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph()
doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('项目名称', '基于OpenHarmony的智能家居设备控制系统'),
    ('赛题编号', 'A9（A组：本科/研究生/高职）'),
    ('出题企业', '苏州未来网络研究院有限公司'),
    ('文档版本', 'v1.0'),
    ('日期', datetime.date.today().strftime('%Y年%m月%d日')),
]
for i, (k, v) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
    for cell in meta_table.rows[i].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
    meta_table.rows[i].cells[0].width = Cm(4)
    meta_table.rows[i].cells[1].width = Cm(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  目录页（手动）
# ══════════════════════════════════════════════════════════════
doc.add_heading('目  录', level=1)
toc_items = [
    '一、项目概述',
    '二、赛题需求分析',
    '三、总体架构设计',
    '四、技术栈选型',
    '五、MQTT 主题设计',
    '六、数据库设计',
    '七、REST API 设计',
    '八、安全三层设计',
    '九、设备联动规则引擎',
    '十、多品牌空调兼容方案',
    '十一、项目文件结构',
    '十二、开发排期',
    '十三、部署方案',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  一、项目概述
# ══════════════════════════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)
doc.add_paragraph(
    '本项目为第十五届"中国软件杯"大学生软件设计大赛 A9 赛题——'
    '基于OpenHarmony操作系统的家居设备控制系统。系统采用云端模拟硬件 + OpenHarmony APP'
    '的架构，将物联网设备（温湿度传感器、人体感应传感器、智能灯、空调、智能门禁）'
    '通过 Python 脚本在阿里云服务器上进行仿真模拟，通过 MQTT 协议与后端服务通信，'
    '最终由 OpenHarmony 原生应用（HAP）提供用户交互界面。'
)
doc.add_heading('核心目标', level=2)
goals = [
    '实现照明中心：远程开关控制 + 人体感应自动开关灯',
    '实现温湿度控制中心：传感器数据采集 + 多品牌空调远程控制（海尔/格力/美的）',
    '实现智能门禁：APP 远程开锁/上锁',
    '实现设备联动：基于规则引擎的自动化场景（人来自动开灯、高温自动开空调等）',
    '实现三层安全加密：传输层 TLS + 数据层 AES + 密钥层 JWT',
    '提交完整作品：HAP 安装包 + 源码 + 设计文档 + PPT + 演示视频',
]
for g in goals:
    add_bullet(doc, g)

# ══════════════════════════════════════════════════════════════
#  二、赛题需求分析
# ══════════════════════════════════════════════════════════════
doc.add_heading('二、赛题需求分析', level=1)

doc.add_heading('2.1 功能需求对照', level=2)
add_table(doc,
    ['赛题要求', '实现方案', '技术路径'],
    [
        ['控制中心：国产主板+网关', '阿里云ECS模拟网关 + Python设备模拟器\n文档说明对应国产主板方案', 'FastAPI + MQTT Broker'],
        ['照明中心：远程开关+人体感应', '灯模拟器订阅MQTT指令 + PIR传感器发布状态\n规则引擎联动', 'MQTT pub/sub + Rule Engine'],
        ['温湿度中心：传感器+多品牌空调', '温湿度模拟器定时上报 + 空调模拟器\n内置海尔/格力/美的指令翻译表', 'Brand Translator Pattern'],
        ['智能门禁：远程开锁/上锁', '门禁模拟器订阅MQTT + 状态反馈\n加密认证码验证', 'MQTT + AES加密'],
        ['协议传输安全', 'TLS 1.3 + AES-256-CBC + JWT\n三层加密独立密钥', 'HTTPS + MQTT over TLS'],
    ]
)

doc.add_heading('2.2 评分标准', level=2)
add_table(doc,
    ['评分项', '分值', '我们的策略'],
    [
        ['功能完整度', '60分', '4大功能全部实现，设备模拟器覆盖所有硬件类型'],
        ['界面美观', '10分', '鸿蒙原生UI，卡片式布局，实时数据可视化'],
        ['可扩展性', '10分', '设备基类抽象，新增设备只需继承BaseDevice；\n规则引擎支持热加载新规则'],
        ['协议安全性', '10分', 'TLS+AES+JWT 三层独立设计，演示时可抓包验证'],
        ['文档质量', '10分', '设计文档+使用手册+部署说明+代码注释完整'],
    ]
)

# ══════════════════════════════════════════════════════════════
#  三、总体架构设计
# ══════════════════════════════════════════════════════════════
doc.add_heading('三、总体架构设计', level=1)

doc.add_heading('3.1 架构分层图', level=2)
doc.add_paragraph(
    '系统采用四层架构：展示层（OpenHarmony APP）→ 云端服务层（FastAPI + 规则引擎 + 安全模块）'
    '→ 消息中间件层（MQTT Broker）→ 设备模拟层（Python 设备模拟器）。'
    '所有组件通过 Docker Compose 在阿里云 ECS 上一键部署。'
)

doc.add_heading('3.2 架构分层说明', level=2)
add_table(doc,
    ['层级', '技术', '职责', '关键组件'],
    [
        ['展示层', 'ArkTS + DevEco Studio\nOpenHarmony SDK', '用户交互界面\n实时数据显示\n设备远程控制', 'DashboardPage\nLightControlPage\nACControlPage\nDoorLockPage'],
        ['云端服务层', 'Python 3.11 + FastAPI\n+ WebSocket', 'REST API\n用户认证\n规则引擎\n数据加密\n历史数据存储', 'auth.py\nrule_engine.py\nsecurity.py\nmqtt_client.py'],
        ['消息中间件', 'Mosquitto MQTT Broker\n端口: 1883/8883(TLS)', '设备消息路由\n发布/订阅管理\nQoS 保障', 'docker-compose 部署\nTLS 证书配置'],
        ['设备模拟层', 'Python asyncio\n+ paho-mqtt', '模拟传感器数据\n响应控制指令\n状态上报', 'temperature_sensor.py\nac_controller.py\ndoor_lock.py 等'],
    ]
)

doc.add_heading('3.3 数据流', level=2)
doc.add_paragraph('上行数据流（传感器 → APP）：')
flows_up = [
    '设备模拟器定时生成传感器数据（温度/湿度/人体感应/设备状态）',
    '通过 paho-mqtt 客户端发布到 MQTT Broker 对应主题',
    'FastAPI 后端订阅 MQTT 主题，收到数据后：存入 SQLite + 触发规则引擎检查',
    '通过 WebSocket 推送给已连接的 OpenHarmony APP',
    'APP 收到实时数据，更新仪表盘/控制页面的 UI 状态',
]
for f in flows_up:
    add_bullet(doc, f)

doc.add_paragraph('下行控制流（APP → 设备）：')
flows_down = [
    '用户在 OpenHarmony APP 上操作（开关灯/调空调/开门）',
    'APP 通过 HTTPS POST 发送指令到 FastAPI 后端',
    '后端验证 JWT 身份 → 加密敏感指令 → 通过 MQTT 发布到设备命令主题',
    '设备模拟器订阅命令主题，收到指令后执行模拟动作',
    '设备模拟器发布响应到状态主题 → 后端更新数据库 → WebSocket 通知 APP 更新 UI',
]
for f in flows_down:
    add_bullet(doc, f)

# ══════════════════════════════════════════════════════════════
#  四、技术栈选型
# ══════════════════════════════════════════════════════════════
doc.add_heading('四、技术栈选型', level=1)
add_table(doc,
    ['层级', '技术选型', '版本', '选型理由'],
    [
        ['APP 开发', 'ArkTS + DevEco Studio', 'OpenHarmony SDK 4.x', '赛题强制要求，提交 HAP 包'],
        ['后端框架', 'FastAPI', '0.111+', '异步高性能，原生 WebSocket 支持，自动生成 API 文档'],
        ['消息队列', 'Mosquitto MQTT', '2.0+', 'IoT 标准协议，轻量级，完善的 TLS 支持'],
        ['数据库', 'SQLite', '3.x', '开发期零配置，单文件存储，可随时迁移至 MySQL'],
        ['设备模拟', 'Python asyncio + paho-mqtt', 'paho 1.6+', '异步并发模拟多设备，官方 MQTT 库'],
        ['加密', 'pycryptodome + PyJWT', '-', 'AES-256 加解密 + JWT 令牌管理'],
        ['部署', 'Docker + Docker Compose', '-', '一键部署，环境隔离，方便评委复现'],
        ['反向代理', 'Nginx', '-', 'HTTPS 终端，反向代理 FastAPI，负载均衡'],
    ]
)

# ══════════════════════════════════════════════════════════════
#  五、MQTT 主题设计
# ══════════════════════════════════════════════════════════════
doc.add_heading('五、MQTT 主题设计', level=1)
doc.add_paragraph(
    'MQTT 主题采用 home/{room_id}/{device_type}/{direction} 的分层结构，'
    '方向分为 sensor（上报）、status（状态）、command（指令）、response（响应）四类。'
)

doc.add_heading('5.1 传感器数据上报主题', level=2)
add_code_block(doc, '''
# 温度传感器 — 每5秒上报一次
主题: home/{room_id}/sensor/temperature
载荷: {"value": 26.5, "unit": "celsius", "device_id": "temp_001", "ts": 1716200000}

# 湿度传感器 — 每5秒上报一次
主题: home/{room_id}/sensor/humidity
载荷: {"value": 65.0, "unit": "percent", "device_id": "hum_001", "ts": 1716200000}

# 人体感应传感器 — 状态变化时上报
主题: home/{room_id}/sensor/pir
载荷: {"presence": true, "device_id": "pir_001", "ts": 1716200000}''')

doc.add_heading('5.2 设备状态上报主题', level=2)
add_code_block(doc, '''
# 灯光状态
主题: home/{room_id}/light/status
载荷: {"power": "on", "brightness": 80, "color": "warm", "device_id": "light_001"}

# 空调状态
主题: home/{room_id}/ac/status
载荷: {"power": "on", "mode": "cool", "temp": 24, "fan": "auto", "brand": "gree",
       "device_id": "ac_001", "ts": 1716200000}

# 门禁状态
主题: home/{room_id}/door/status
载荷: {"locked": true, "device_id": "door_001", "ts": 1716200000}''')

doc.add_heading('5.3 控制指令主题', level=2)
add_code_block(doc, '''
# 灯控指令
主题: home/{room_id}/light/command
载荷: {"action": "on", "brightness": 70, "color": "warm"}
     {"action": "off"}

# 空调指令（统一指令模型，由后端翻译为品牌协议）
主题: home/{room_id}/ac/command
载荷: {"action": "set", "mode": "cool", "temp": 24, "fan": "auto"}
     {"action": "off"}

# 门禁指令（含加密认证码）
主题: home/{room_id}/door/command
载荷: {"action": "unlock", "auth_code": "<AES加密的一次性认证码>"}
     {"action": "lock"}''')

doc.add_heading('5.4 控制响应主题', level=2)
add_code_block(doc, '''
主题: home/{room_id}/light/response
载荷: {"success": true, "state": {"power": "on", "brightness": 70}}

主题: home/{room_id}/ac/response
载荷: {"success": true, "state": {"power": "on", "mode": "cool", "temp": 24}}

主题: home/{room_id}/door/response
载荷: {"success": true, "state": {"locked": false}}''')

# ══════════════════════════════════════════════════════════════
#  六、数据库设计
# ══════════════════════════════════════════════════════════════
doc.add_heading('六、数据库设计', level=1)

doc.add_heading('6.1 ER 关系', level=2)
doc.add_paragraph(
    'User (1) ──< (N) Room (1) ──< (N) Device\n'
    'Device (1) ──< (N) SensorData\n'
    'Device (1) ──< (N) DeviceLog\n'
    'AutomationRule 独立实体，通过 JSON 字段关联设备 ID'
)

doc.add_heading('6.2 建表 SQL', level=2)
add_code_block(doc, '''
-- 用户表
CREATE TABLE users (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    username    TEXT NOT NULL UNIQUE,
    password_hash TEXT NOT NULL,
    role        TEXT DEFAULT 'user',       -- user / admin
    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
);

-- 房间表
CREATE TABLE rooms (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    name        TEXT NOT NULL,             -- 如 '客厅', '卧室', '书房'
    floor       INTEGER DEFAULT 1,
    description TEXT,
    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
);

-- 设备表
CREATE TABLE devices (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    room_id     INTEGER NOT NULL REFERENCES rooms(id),
    type        TEXT NOT NULL,             -- temperature_sensor | humidity_sensor
                                           -- pir_sensor | light | ac | door_lock
    name        TEXT NOT NULL,             -- '客厅温度传感器'
    brand       TEXT,                      -- 仅 ac 类型使用: haier/gree/midea/generic
    mqtt_topic  TEXT NOT NULL,             -- MQTT 主题前缀
    status_json TEXT DEFAULT '{}',          -- 当前状态 JSON
    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
);

-- 传感器数据表（历史记录）
CREATE TABLE sensor_data (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    device_id   INTEGER NOT NULL REFERENCES devices(id),
    data_type   TEXT NOT NULL,             -- temperature | humidity | presence | status
    value       REAL,                      -- 数值
    extra_json  TEXT DEFAULT '{}',          -- 扩展字段
    timestamp   DATETIME DEFAULT CURRENT_TIMESTAMP
);
CREATE INDEX idx_sensor_data_device_ts ON sensor_data(device_id, timestamp);

-- 设备操作日志表
CREATE TABLE device_log (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    device_id   INTEGER NOT NULL REFERENCES devices(id),
    action      TEXT NOT NULL,             -- 'turn_on', 'set_temp', 'unlock' 等
    detail      TEXT,                      -- 操作详情
    user_id     INTEGER REFERENCES users(id),
    timestamp   DATETIME DEFAULT CURRENT_TIMESTAMP
);

-- 联动规则表
CREATE TABLE automation_rules (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    name        TEXT NOT NULL,             -- '人来开灯'
    condition_json TEXT NOT NULL,           -- 触发条件
    action_json TEXT NOT NULL,              -- 执行动作
    enabled     INTEGER DEFAULT 1,
    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
);''')

doc.add_heading('6.3 初始数据示例', level=2)
add_code_block(doc, '''
-- 客厅
INSERT INTO rooms VALUES (1, '客厅', 1, '家庭客厅区域', datetime('now'));
INSERT INTO devices VALUES (1, 1, 'temperature_sensor', '客厅温度', NULL,
    'home/livingroom/sensor/temperature', '{}', datetime('now'));
INSERT INTO devices VALUES (2, 1, 'humidity_sensor', '客厅湿度', NULL,
    'home/livingroom/sensor/humidity', '{}', datetime('now'));
INSERT INTO devices VALUES (3, 1, 'pir_sensor', '客厅人体感应', NULL,
    'home/livingroom/sensor/pir', '{}', datetime('now'));
INSERT INTO devices VALUES (4, 1, 'light', '客厅主灯', NULL,
    'home/livingroom/light', '{"power":"off","brightness":0}', datetime('now'));
INSERT INTO devices VALUES (5, 1, 'ac', '客厅空调', 'gree',
    'home/livingroom/ac', '{"power":"off","mode":"cool","temp":26}', datetime('now'));

-- 联动规则
INSERT INTO automation_rules VALUES (1, '人来开灯',
    '{"trigger":"pir_sensor","field":"presence","operator":"eq","value":true,
      "and":[{"trigger":"light","field":"power","operator":"eq","value":"off"}]}',
    '[{"device_type":"light","action":"on","params":{"brightness":80}}]',
    1, datetime('now'));

INSERT INTO automation_rules VALUES (2, '人走关灯',
    '{"trigger":"pir_sensor","field":"presence","operator":"eq","value":false,
      "and":[{"trigger":"light","field":"power","operator":"eq","value":"on"},
             {"trigger":"light","field":"on_duration_sec","operator":"gt","value":300}]}',
    '[{"device_type":"light","action":"off","params":{}}]',
    1, datetime('now'));

INSERT INTO automation_rules VALUES (3, '高温自动制冷',
    '{"trigger":"temperature_sensor","field":"value","operator":"gt","value":28,
      "and":[{"trigger":"ac","field":"power","operator":"eq","value":"off"}]}',
    '[{"device_type":"ac","action":"set","params":{"power":"on","mode":"cool","temp":26}}]',
    1, datetime('now'));

INSERT INTO automation_rules VALUES (4, '高湿自动除湿',
    '{"trigger":"humidity_sensor","field":"value","operator":"gt","value":80,
      "and":[{"trigger":"ac","field":"power","operator":"eq","value":"off"}]}',
    '[{"device_type":"ac","action":"set","params":{"power":"on","mode":"dehumidify"}}]',
    1, datetime('now'));''')

# ══════════════════════════════════════════════════════════════
#  七、REST API 设计
# ══════════════════════════════════════════════════════════════
doc.add_heading('七、REST API 设计', level=1)

doc.add_heading('7.1 认证模块', level=2)
add_table(doc,
    ['方法', '路径', '说明', '认证'],
    [
        ['POST', '/api/auth/login', '用户登录，返回 JWT token', '否'],
        ['POST', '/api/auth/register', '用户注册', '否'],
        ['GET', '/api/auth/me', '获取当前用户信息', 'JWT'],
    ]
)

doc.add_heading('7.2 房间管理', level=2)
add_table(doc,
    ['方法', '路径', '说明', '认证'],
    [
        ['GET', '/api/rooms', '获取房间列表（含设备数量）', 'JWT'],
        ['GET', '/api/rooms/{id}', '获取房间详情（含设备实时状态）', 'JWT'],
        ['POST', '/api/rooms', '添加房间', 'JWT'],
        ['PUT', '/api/rooms/{id}', '修改房间信息', 'JWT'],
        ['DELETE', '/api/rooms/{id}', '删除房间', 'JWT'],
    ]
)

doc.add_heading('7.3 设备管理', level=2)
add_table(doc,
    ['方法', '路径', '说明', '认证'],
    [
        ['GET', '/api/devices', '设备列表（可按 room_id/type 筛选）', 'JWT'],
        ['GET', '/api/devices/{id}', '设备详情 + 当前状态', 'JWT'],
        ['POST', '/api/devices/{id}/command', '发送控制指令', 'JWT'],
        ['PUT', '/api/devices/{id}', '修改设备信息', 'JWT'],
    ]
)

doc.add_heading('7.4 历史数据', level=2)
add_table(doc,
    ['方法', '路径', '说明', '认证'],
    [
        ['GET', '/api/data/sensors', '传感器历史数据\n?device_id=&type=&start=&end=&limit=', 'JWT'],
        ['GET', '/api/data/logs', '设备操作日志\n?device_id=&user_id=&start=&end=', 'JWT'],
    ]
)

doc.add_heading('7.5 联动规则', level=2)
add_table(doc,
    ['方法', '路径', '说明', '认证'],
    [
        ['GET', '/api/rules', '规则列表', 'JWT'],
        ['POST', '/api/rules', '创建规则', 'JWT'],
        ['PUT', '/api/rules/{id}', '更新规则', 'JWT'],
        ['DELETE', '/api/rules/{id}', '删除规则', 'JWT'],
        ['POST', '/api/rules/{id}/toggle', '启用/禁用规则', 'JWT'],
    ]
)

doc.add_heading('7.6 WebSocket 实时数据', level=2)
add_code_block(doc, '''
连接: ws://<server>:8000/ws/realtime?token=<jwt_token>

服务端推送格式:
{
    "type": "sensor_update",
    "device_id": 4,
    "room_id": 1,
    "data": {"temperature": 26.5, "humidity": 65.0, "presence": true},
    "timestamp": 1716200000
}

{
    "type": "device_state",
    "device_id": 5,
    "state": {"power": "on", "mode": "cool", "temp": 24},
    "timestamp": 1716200000
}

{
    "type": "rule_triggered",
    "rule_id": 1,
    "rule_name": "人来开灯",
    "result": {"device_id": 4, "action": "on", "success": true},
    "timestamp": 1716200000
}''')

# ══════════════════════════════════════════════════════════════
#  八、安全三层设计
# ══════════════════════════════════════════════════════════════
doc.add_heading('八、安全三层设计', level=1)
doc.add_paragraph(
    '对应评分标准中"协议传输安全性"10 分，系统从传输层、数据层、密钥层三个维度'
    '独立设计安全方案，层层递进，每层不共享密钥。'
)

add_table(doc,
    ['安全层', '技术方案', '密钥来源', '保护范围'],
    [
        ['传输层', 'HTTPS (TLS 1.3)\nMQTT over TLS (端口8883)', 'CA 签发的 SSL/TLS 证书\n（自签名证书用于开发）', 'APP ↔ 后端 全链路加密\n后端 ↔ MQTT Broker 加密'],
        ['数据层', 'AES-256-CBC 对称加密\n敏感字段载荷加密', '由密钥层 JWT payload 中的\naes_key 字段提供', '门禁认证码、空调控制指令\n等敏感操作载荷'],
        ['密钥层', 'JWT (HS256) 身份令牌\n用户独立 AES 密钥\n密钥定期轮换', '服务端密钥 master_secret\n每个用户的 aes_key 独立生成', '用户身份认证\nAES 密钥安全分发'],
    ]
)

doc.add_heading('8.1 加密流程示意', level=2)
add_code_block(doc, '''
【APP 开门禁流程】
1. APP 请求开门 → 携带 JWT token
2. 后端验证 JWT → 提取用户 aes_key
3. 后端生成一次性认证码 → AES-256-CBC(auth_code, aes_key)
4. 后端通过 MQTT over TLS 发布加密后的 auth_code 到门禁指令主题
5. 门禁模拟器解密验证 → 执行 → 上锁/开锁

【密钥轮换策略】
- 用户 AES 密钥在登录时生成，token 过期（24h）后自动轮换
- 旧密钥保留 1 小时用于解密仍在传输中的消息
- 服务端 master_secret 通过环境变量注入，不写入代码''')

# ══════════════════════════════════════════════════════════════
#  九、设备联动规则引擎
# ══════════════════════════════════════════════════════════════
doc.add_heading('九、设备联动规则引擎', level=1)
doc.add_paragraph(
    '规则引擎是系统的智能核心，监听 MQTT 传感器数据流，当条件满足时自动触发设备动作。'
    '支持热加载（通过 API 增删改规则后立即生效，无需重启服务）。'
)

doc.add_heading('9.1 引擎架构', level=2)
add_code_block(doc, '''
RuleEngine (单例)
├── _rules: List[Rule]           # 内存中的规则列表
├── _device_states: Dict         # 设备当前状态缓存
├── reload_rules()               # 从数据库重新加载规则
├── on_sensor_data(topic, payload) # 收到传感器数据时触发
│   ├── 更新 _device_states 缓存
│   ├── 遍历所有启用的规则
│   ├── evaluate(rule.condition) → True/False
│   └── execute(rule.actions)    → 通过 MQTT 下发指令
└── execute_actions(actions)     # 执行动作列表''')

doc.add_heading('9.2 规则定义格式', level=2)
add_code_block(doc, '''
# 规则条件 (condition_json)
{
    "trigger": "temperature_sensor",    # 触发源设备类型
    "field": "value",                   # 比较字段
    "operator": "gt",                   # gt | lt | eq | neq | changed
    "value": 28,                        # 阈值
    "room_id": null,                    # null 表示全局，指定则只在该房间生效
    "and": [                            # 可选：附加条件（全部满足）
        {"trigger": "ac", "field": "power", "operator": "eq", "value": "off"}
    ]
}

# 规则动作 (action_json) — 支持多动作顺序执行
[
    {
        "device_type": "ac",            # 目标设备类型
        "room_id": "same",              # same = 与触发源同一房间 | 具体房间ID
        "action": "set",                # 动作名
        "params": {                     # 动作参数
            "power": "on",
            "mode": "cool",
            "temp": 26
        }
    }
]''')

# ══════════════════════════════════════════════════════════════
#  十、多品牌空调兼容方案
# ══════════════════════════════════════════════════════════════
doc.add_heading('十、多品牌空调兼容方案', level=1)
doc.add_paragraph(
    '赛题明确要求兼容海尔、格力、美的三个品牌。本方案采用"统一指令模型 + 品牌翻译器"'
    '设计模式。APP 和后端只操作统一指令模型，由设备模拟器内部的品牌翻译器'
    '将统一指令转换为各品牌特有的控制协议。'
)

doc.add_heading('10.1 统一指令模型', level=2)
add_code_block(doc, '''
# 后端和 APP 统一使用的指令格式
{
    "power": "on" | "off",
    "mode": "cool" | "heat" | "dehumidify" | "fan_only" | "auto",
    "temp": 16-30,              # 目标温度（摄氏度）
    "fan": "auto" | "low" | "medium" | "high",
    "swing": "on" | "off"       # 摆风
}''')

doc.add_heading('10.2 品牌翻译器', level=2)
add_code_block(doc, '''
# ac_brand.py — 品牌指令翻译表

UNIVERSAL_TO_BRAND = {
    "gree": {
        "power":     {"on": "PWR_ON",  "off": "PWR_OFF"},
        "mode":      {"cool": "MODE_COOL", "heat": "MODE_HEAT",
                      "dehumidify": "MODE_DRY", "fan_only": "MODE_FAN",
                      "auto": "MODE_AUTO"},
        "fan":       {"auto": "FAN_AUTO", "low": "FAN_1",
                      "medium": "FAN_2", "high": "FAN_3"},
        "temp":      lambda t: f"TEMP_{t}",      # TEMP_24
    },
    "haier": {
        "power":     {"on": "POWER=1", "off": "POWER=0"},
        "mode":      {"cool": "MODE=COOLING", "heat": "MODE=HEATING",
                      "dehumidify": "MODE=DRY", "fan_only": "MODE=FAN",
                      "auto": "MODE=SMART"},
        "fan":       {"auto": "FAN=AUTO", "low": "FAN=LOW",
                      "medium": "FAN=MED", "high": "FAN=HIGH"},
        "temp":      lambda t: f"SET_TEMP={t}",  # SET_TEMP=24
    },
    "midea": {
        "power":     {"on": 1,  "off": 0},
        "mode":      {"cool": 2, "heat": 3, "dehumidify": 4,
                      "fan_only": 1, "auto": 0},
        "fan":       {"auto": 1024, "low": 40, "medium": 60, "high": 80},
        "temp":      lambda t: t,                  # 直接数值 24
    },
}''')

doc.add_heading('10.3 APP 中的品牌切换体验', level=2)
doc.add_paragraph(
    '在 APP 空调控制页面，用户可以从下拉框中选择空调品牌（海尔/格力/美的/通用）。'
    '切换品牌后，APP 发送同样的控制指令，后端自动路由到对应品牌的空调模拟器，'
    '模拟器内部使用品牌翻译器转换指令。用户无感知，但通过设备状态返回可以看到'
    '不同品牌的响应格式差异，体现系统的多品牌兼容能力。'
)

# ══════════════════════════════════════════════════════════════
#  十一、项目文件结构
# ══════════════════════════════════════════════════════════════
doc.add_heading('十一、项目文件结构', level=1)
add_code_block(doc, '''
smart-home-A9/
├── cloud/                              # 阿里云服务器端（全部 Python）
│   ├── backend/                        # FastAPI 后端服务
│   │   ├── app/
│   │   │   ├── __init__.py
│   │   │   ├── main.py                 # FastAPI 入口 + 生命周期管理
│   │   │   ├── config.py               # 配置（数据库/MQTT/JWT密钥）
│   │   │   ├── api/
│   │   │   │   ├── __init__.py
│   │   │   │   ├── auth.py             # POST /api/auth/login, /register
│   │   │   │   ├── rooms.py            # CRUD /api/rooms
│   │   │   │   ├── devices.py          # GET /api/devices + POST command
│   │   │   │   ├── data.py             # GET /api/data/sensors, /logs
│   │   │   │   └── rules.py            # CRUD /api/rules
│   │   │   ├── models/
│   │   │   │   ├── __init__.py
│   │   │   │   ├── user.py
│   │   │   │   ├── device.py
│   │   │   │   ├── room.py
│   │   │   │   └── sensor_data.py
│   │   │   ├── services/
│   │   │   │   ├── __init__.py
│   │   │   │   ├── mqtt_client.py      # MQTT 连接管理 + 主题订阅
│   │   │   │   ├── rule_engine.py      # 联动规则引擎（核心）
│   │   │   │   ├── security.py         # JWT + AES 加密
│   │   │   │   └── ac_brand.py         # 多品牌空调指令翻译
│   │   │   └── database/
│   │   │       ├── __init__.py
│   │   │       ├── connection.py       # SQLite 连接管理
│   │   │       └── init_db.py          # 建表 + 初始数据
│   │   ├── requirements.txt
│   │   └── Dockerfile
│   ├── simulators/                     # 硬件设备模拟器
│   │   ├── __init__.py
│   │   ├── base_device.py              # 设备基类（抽象类）
│   │   ├── temperature_sensor.py       # 温度传感器（正态分布模拟）
│   │   ├── humidity_sensor.py          # 湿度传感器
│   │   ├── pir_sensor.py               # 人体感应传感器
│   │   ├── light_controller.py         # 灯控制器（支持亮度/色温）
│   │   ├── ac_controller.py            # 空调控制器（多品牌）
│   │   ├── door_lock.py                # 门禁控制器
│   │   └── simulator_manager.py        # 统一启停管理
│   └── docker-compose.yml              # 一键部署编排
├── openharmony/                        # 鸿蒙 APP 源码
│   └── entry/src/main/ets/
│       ├── entryability/EntryAbility.ets
│       ├── pages/
│       │   ├── Index.ets               # 启动/引导页
│       │   ├── LoginPage.ets           # 登录页
│       │   ├── DashboardPage.ets       # 仪表盘（总览+实时数据）
│       │   ├── LightControlPage.ets    # 照明控制
│       │   ├── ACControlPage.ets       # 空调控制（含品牌切换）
│       │   └── DoorLockPage.ets        # 门禁控制
│       ├── model/
│       │   └── DeviceModel.ets         # 数据模型定义
│       └── common/
│           ├── ApiClient.ets           # HTTP 请求封装
│           └── MqttClient.ets          # MQTT 客户端封装
└── docs/                               # 提交文档
    ├── 产品设计文档.md
    ├── 使用手册.md
    ├── 部署说明.md
    └── PPT大纲.md''')

# ══════════════════════════════════════════════════════════════
#  十二、开发排期
# ══════════════════════════════════════════════════════════════
doc.add_heading('十二、开发排期（共 5 周）', level=1)

add_table(doc,
    ['周次', '日期', '阶段', '任务内容', '产出物'],
    [
        ['W1', '5/20-5/26', '基础设施',
         '① 阿里云 ECS 申请+环境配置\n② Python 项目骨架搭建\n③ 数据库建表+初始数据\n④ FastAPI 启动+认证模块\n⑤ Mosquitto MQTT Broker 部署',
         '可运行的后端骨架\n用户能登录注册\nMQTT 服务器通'],
        ['W2', '5/27-6/2', '设备模拟',
         '① BaseDevice 基类设计\n② 温湿度传感器模拟器\n③ PIR 人体感应模拟器\n④ 灯/空调/门禁控制器\n⑤ 设备管理 API\n⑥ WebSocket 实时推送',
         '全部 6 种设备模拟器\n能通过 API 查看设备状态\nAPP 能收到实时推送'],
        ['W3', '6/3-6/9', '核心逻辑',
         '① 规则引擎开发\n② 安全模块：TLS+AES+JWT\n③ 多品牌空调翻译器\n④ 鸿蒙 APP 项目创建\n⑤ 登录页+仪表盘页',
         '规则引擎可工作\n三层加密完成\n鸿蒙 APP 能登录'],
        ['W4', '6/10-6/16', 'APP+联调',
         '① 照明控制页+空调控制页+门禁页\n② 前后端全链路联调\n③ 品牌切换功能验证\n④ 规则引擎端到端测试\n⑤ UI 优化美化',
         'APP 全部页面完成\n前后端联调通过'],
        ['W5', '6/17-6/23', '收尾',
         '① 全系统集成测试\n② Bug 修复\n③ 录制演示视频（≤7分钟）\n④ 编写产品设计文档\n⑤ 编写使用手册\n⑥ 制作答辩 PPT',
         '作品完整可提交'],
        ['缓冲', '6/24-6/30', '提交',
         '① 最后修 bug\n② 打包 HAP\n③ 整理提交材料\n④ 6月30日15:00前提交',
         '提交完成'],
    ]
)

# ══════════════════════════════════════════════════════════════
#  十三、部署方案
# ══════════════════════════════════════════════════════════════
doc.add_heading('十三、部署方案', level=1)

doc.add_heading('13.1 阿里云 ECS 配置建议', level=2)
add_table(doc,
    ['配置项', '建议值', '说明'],
    [
        ['实例规格', '2 vCPU / 4 GB 内存', 'ecs.c7.large 或 ecs.g7.large'],
        ['操作系统', 'Ubuntu 22.04 LTS', 'Python 3.11 原生支持好'],
        ['系统盘', '40 GB ESSD', '系统和 Docker 镜像足够'],
        ['带宽', '按量计费 5 Mbps', '开发期够用，后期可调'],
        ['安全组', '开放 22/80/443/8000/8883', 'SSH/HTTP/HTTPS/API/MQTT TLS'],
    ]
)

doc.add_heading('13.2 Docker Compose 一键部署', level=2)
add_code_block(doc, '''
# docker-compose.yml
version: '3.8'
services:
  mqtt:
    image: eclipse-mosquitto:2
    ports:
      - "1883:1883"
      - "8883:8883"
    volumes:
      - ./mosquitto/config:/mosquitto/config
      - ./mosquitto/data:/mosquitto/data
      - ./mosquitto/certs:/mosquitto/certs
    restart: unless-stopped

  backend:
    build: ./backend
    ports:
      - "8000:8000"
    environment:
      - DATABASE_URL=sqlite:///data/smart_home.db
      - MQTT_BROKER=mqtt
      - MQTT_PORT=1883
      - JWT_SECRET=${JWT_SECRET}
      - JWT_EXPIRE_HOURS=24
    volumes:
      - ./data:/app/data
    depends_on:
      - mqtt
    restart: unless-stopped

  simulators:
    build: ./simulators
    environment:
      - MQTT_BROKER=mqtt
      - MQTT_PORT=1883
    depends_on:
      - mqtt
    restart: unless-stopped

  nginx:
    image: nginx:alpine
    ports:
      - "80:80"
      - "443:443"
    volumes:
      - ./nginx/nginx.conf:/etc/nginx/nginx.conf
      - ./nginx/certs:/etc/nginx/certs
    depends_on:
      - backend
    restart: unless-stopped''')

doc.add_heading('13.3 启动命令', level=2)
add_code_block(doc, '''
# 1. 上传代码到阿里云 ECS
scp -r ./smart-home-A9 root@<ECS公网IP>:/opt/

# 2. SSH 登录后启动
cd /opt/smart-home-A9/cloud
docker compose up -d

# 3. 验证
curl http://localhost:8000/docs          # FastAPI 自动文档
mosquitto_sub -h localhost -t 'home/#'   # 查看 MQTT 消息流''')

# ══════════════════════════════════════════════════════════════
#  附录：关键风险与应对
# ══════════════════════════════════════════════════════════════
doc.add_heading('附录：关键风险与应对', level=1)
add_table(doc,
    ['风险', '等级', '应对措施'],
    [
        ['OpenHarmony 开发学习曲线陡', '高', '先做 Web 前端验证所有接口，再迁移到 ArkTS；\n使用 DevEco Studio 模板加速开发'],
        ['时间紧张（5周）', '高', '严格按排期执行，每周日检查进度；\n砍掉非核心功能，优先保证评分项'],
        ['阿里云费用', '低', '使用学生优惠/免费试用套餐；\n开发期用本地环境，提交前再迁移云上'],
        ['MQTT + TLS 配置复杂', '中', '先用非 TLS 开发调试，最后一周再加 TLS；\n参考 Mosquitto 官方 TLS 配置文档'],
        ['多品牌空调协议差异', '低', '只需模拟指令格式差异，不需要对接真实空调；\n翻译器模式改动成本低'],
    ]
)

# ── 保存 ──
output_path = 'C:/Mycode/smart-home-A9/A9智能家居项目框架文档.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
